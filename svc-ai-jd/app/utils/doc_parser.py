import os
import logging

logger = logging.getLogger(__name__)


def load_file_content(file_path: str, parse_mode: str = "auto") -> str:
    """
    Detects the file extension and returns parsed content from:
    - Text (.txt)
    - PDF (.pdf) — 3-tier: LlamaParse → Docling → PyMuPDF
    - Word Document (.docx) — 2-tier: Docling → python-docx

    parse_mode:
        "auto"    — try tiers in order, silent fallback (default)
        "quality" — same + log warnings on each fallback
        "fast"    — PyMuPDF/python-docx only, skip cloud/heavy parsers

    Raises:
        FileNotFoundError or ValueError if file is invalid or unsupported.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # Allow env override for parse_mode default
    if parse_mode == "auto":
        parse_mode = os.getenv("DOC_PARSE_MODE", "auto")

    _, ext = os.path.splitext(file_path.lower())

    if ext == ".txt":
        return _load_txt(file_path)
    elif ext == ".pdf":
        return _load_pdf(file_path, parse_mode)
    elif ext == ".docx":
        return _load_docx(file_path, parse_mode)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _load_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


# ---------------------------------------------------------------------------
# PDF: 3-tier parsing — LlamaParse → Docling → PyMuPDF
# ---------------------------------------------------------------------------

def _load_pdf(path: str, parse_mode: str = "auto") -> str:
    if parse_mode != "fast":
        # Tier 1: LlamaParse (cloud)
        result = _try_llamaparse_pdf(path)
        if result is not None:
            logger.info("PDF parsed with LlamaParse (Tier 1)")
            return result + _extract_pdf_hyperlinks_suffix(path, result)
        if parse_mode == "quality":
            logger.warning("LlamaParse unavailable, falling back to Docling")

        # Tier 2: Docling (local)
        result = _try_docling_pdf(path)
        if result is not None:
            logger.info("PDF parsed with Docling (Tier 2)")
            return result + _extract_pdf_hyperlinks_suffix(path, result)
        if parse_mode == "quality":
            logger.warning("Docling unavailable, falling back to PyMuPDF")

    # Tier 3: PyMuPDF (always available)
    logger.info("PDF parsed with PyMuPDF (Tier 3)")
    return _load_pdf_pymupdf(path)


def _try_llamaparse_pdf(path: str) -> str | None:
    api_key = os.getenv("LLAMA_PARSE_API_KEY", "").strip()
    if not api_key:
        return None
    try:
        from llama_parse import LlamaParse
        parser = LlamaParse(api_key=api_key, result_type="markdown")
        documents = parser.load_data(path)
        if documents:
            return "\n\n".join(doc.text for doc in documents)
        return None
    except ImportError:
        logger.debug("llama-parse not installed, skipping Tier 1")
        return None
    except Exception as e:
        logger.warning(f"LlamaParse failed: {e}")
        return None


def _try_docling_pdf(path: str) -> str | None:
    try:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(path)
        md = result.document.export_to_markdown()
        if md and md.strip():
            return md
        return None
    except ImportError:
        logger.debug("docling not installed, skipping Tier 2")
        return None
    except Exception as e:
        logger.warning(f"Docling failed: {e}")
        return None


def _load_pdf_pymupdf(path: str) -> str:
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run `pip install pymupdf`.")

    doc = fitz.open(path)
    text = ""
    all_link_urls = []
    for page in doc:
        text += page.get_text()
        for link in page.get_links():
            if link.get("kind") == 2 and link.get("uri"):
                uri = link["uri"]
                if uri not in text and uri not in all_link_urls:
                    all_link_urls.append(uri)
    doc.close()

    if all_link_urls:
        text += "\n\n--- Embedded Hyperlinks ---\n"
        text += "\n".join(all_link_urls)

    return text


def _extract_pdf_hyperlinks_suffix(path: str, existing_text: str) -> str:
    """Extract PDF annotation hyperlinks not already present in the parsed text."""
    try:
        import fitz
    except ImportError:
        return ""

    try:
        doc = fitz.open(path)
        urls = []
        for page in doc:
            for link in page.get_links():
                if link.get("kind") == 2 and link.get("uri"):
                    uri = link["uri"]
                    if uri not in existing_text and uri not in urls:
                        urls.append(uri)
        doc.close()

        if urls:
            return "\n\n--- Embedded Hyperlinks ---\n" + "\n".join(urls)
    except Exception as e:
        logger.debug(f"Hyperlink extraction failed: {e}")

    return ""


# ---------------------------------------------------------------------------
# DOCX: 2-tier parsing — Docling → python-docx
# ---------------------------------------------------------------------------

def _load_docx(path: str, parse_mode: str = "auto") -> str:
    if parse_mode != "fast":
        # Tier 2: Docling (local)
        result = _try_docling_docx(path)
        if result is not None:
            logger.info("DOCX parsed with Docling (Tier 2)")
            return result + _extract_docx_hyperlinks_suffix(path, result)
        if parse_mode == "quality":
            logger.warning("Docling unavailable for DOCX, falling back to python-docx")

    # Tier 3: python-docx (always available)
    logger.info("DOCX parsed with python-docx (Tier 3)")
    return _load_docx_python_docx(path)


def _try_docling_docx(path: str) -> str | None:
    try:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(path)
        md = result.document.export_to_markdown()
        if md and md.strip():
            return md
        return None
    except ImportError:
        logger.debug("docling not installed, skipping Tier 2 for DOCX")
        return None
    except Exception as e:
        logger.warning(f"Docling DOCX failed: {e}")
        return None


def _load_docx_python_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx not installed. Run `pip install python-docx`.")

    doc = docx.Document(path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    full_text = "\n".join(paragraphs)

    link_urls = []
    for para in doc.paragraphs:
        for rel in para.part.rels.values():
            if "hyperlink" in rel.reltype:
                url = rel._target
                if url and url not in full_text and url not in link_urls:
                    link_urls.append(url)

    if link_urls:
        full_text += "\n\n--- Embedded Hyperlinks ---\n"
        full_text += "\n".join(link_urls)

    return full_text


def _extract_docx_hyperlinks_suffix(path: str, existing_text: str) -> str:
    """Extract DOCX hyperlinks not already present in the parsed text."""
    try:
        import docx
    except ImportError:
        return ""

    try:
        doc = docx.Document(path)
        urls = []
        for para in doc.paragraphs:
            for rel in para.part.rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel._target
                    if url and url not in existing_text and url not in urls:
                        urls.append(url)

        if urls:
            return "\n\n--- Embedded Hyperlinks ---\n" + "\n".join(urls)
    except Exception as e:
        logger.debug(f"DOCX hyperlink extraction failed: {e}")

    return ""
